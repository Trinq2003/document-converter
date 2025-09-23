import argparse
import os
from datetime import datetime
import platform
from concurrent.futures import ProcessPoolExecutor, as_completed
from typing import List, Tuple
import multiprocessing

try:
    from tqdm import tqdm
    TQDM_AVAILABLE = True
except ImportError:
    TQDM_AVAILABLE = False
    # Create dummy tqdm function
    def tqdm(iterable=None, **kwargs):
        return iterable if iterable is not None else iter([])

# Windows-specific imports - only available on Windows
try:
    import pythoncom
    import win32com.client as win32
    from win32com.client import constants as c
    WINDOWS_AVAILABLE = True
    print(f"Windows available: {WINDOWS_AVAILABLE}")
except ImportError:
    WINDOWS_AVAILABLE = False
    # Create dummy objects for type hints when not on Windows
    pythoncom = None
    win32 = None
    c = None


WATERMARK_TAG = "AI_RACE_WATERMARK"


def rgb(r: int, g: int, b: int) -> int:
    """Convert RGB tuple to VB color integer."""
    return (b << 16) + (g << 8) + r


def clear_existing_watermarks(header_shapes) -> None:
    """Remove shapes previously added by this script (by alt text tag)."""
    # Word collections are 1-based; iterate backwards when deleting
    for i in range(header_shapes.Count, 0, -1):
        shape = header_shapes.Item(i)
        try:
            alt_text = getattr(shape, "AlternativeText", "") or ""
            if alt_text.startswith(WATERMARK_TAG):
                shape.Delete()
        except Exception:
            # Ignore shapes that don't expose expected properties
            continue


def add_tiled_watermarks_to_header(header, text: str) -> None:
    """Tile four watermark variants across the page in the section header."""
    shapes = header.Shapes
    clear_existing_watermarks(shapes)

    # Some Word object models may not expose Range.Sections(1) here; fall back to
    # header.PageSetup when available.
    try:
        section = header.Range.Sections(1)
        page_width = float(section.PageSetup.PageWidth)
        page_height = float(section.PageSetup.PageHeight)
    except Exception:
        page_width = float(header.PageSetup.PageWidth)
        page_height = float(header.PageSetup.PageHeight)

    # Four variants: different sizes and rotations
    variants = [
        {"font_size": 40, "rotation": 315},
        {"font_size": 16, "rotation": 330},
        {"font_size": 12, "rotation": 345},
        {"font_size": 30, "rotation": 300},
    ]

    # Spacing between tiles (points). Tuned for legible tiling across common page sizes.
    step_x = 320.0
    step_y = 240.0

    y = -120.0
    row_index = 0
    tile_index = 0

    # Extend slightly beyond page bounds so tiles reach borders when rotated
    while y <= page_height + 120.0:
        x_offset = -160.0 + (step_x / 2.0 if row_index % 2 == 1 else 0.0)
        x = x_offset
        while x <= page_width + 160.0:
            variant = variants[tile_index % len(variants)]
            font_size = variant["font_size"]
            rotation = variant["rotation"]

            # Fallbacks for Office enum constants if not generated
            msoTextEffect1 = getattr(c, "msoTextEffect1", 0)
            msoSendBehindText = getattr(c, "msoSendBehindText", 5)
            wdRelHPage = getattr(c, "wdRelativeHorizontalPositionPage", 1)
            wdRelVPage = getattr(c, "wdRelativeVerticalPositionPage", 1)

            shape = shapes.AddTextEffect(
                msoTextEffect1,
                text,
                "Arial",
                font_size,
                False,
                False,
                x,
                y,
            )

            # Visual styling: light gray, high transparency, behind content
            shape.Rotation = rotation
            shape.Line.Visible = False
            shape.Fill.Visible = True
            shape.Fill.ForeColor.RGB = rgb(180, 180, 180)
            shape.Fill.Transparency = 0.5
            shape.WrapFormat.AllowOverlap = True
            shape.RelativeHorizontalPosition = wdRelHPage
            shape.RelativeVerticalPosition = wdRelVPage
            try:
                shape.LockAspectRatio = True
            except Exception:
                pass
            try:
                shape.ZOrder(msoSendBehindText)
            except Exception:
                pass

            # Tag for idempotent reruns
            try:
                shape.AlternativeText = f"{WATERMARK_TAG}::{text}"
            except Exception:
                pass

            x += step_x
            tile_index += 1

        y += step_y
        row_index += 1


def add_watermarks_to_docx(input_path: str, output_path: str) -> None:
    """Open a .docx, add complex watermarks in headers of all sections, save copy."""
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Input file not found: {input_path}")

    # Try Windows COM method first (preferred for complex watermarks)
    if WINDOWS_AVAILABLE:
        return _add_watermarks_windows(input_path, output_path)
    else:
        # Fallback to basic watermarking for Linux
        try:
            return _add_watermarks_basic(input_path, output_path)
        except ImportError as e:
            raise OSError(
                f"Watermark functionality requires dependencies. "
                f"For Windows: install with 'pip install -e .[windows]'. "
                f"For Linux: install python-docx with 'pip install python-docx'. "
                f"Original error: {e}"
            )


def process_single_file(args: Tuple[str, str, str]) -> Tuple[str, bool, str]:
    """
    Process a single file for parallel execution.

    Args:
        args: Tuple of (input_path, output_path, method)

    Returns:
        Tuple of (input_path, success, error_message)
    """
    input_path, output_path, method = args

    try:
        # Ensure output directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        if method == "windows" and WINDOWS_AVAILABLE:
            _add_watermarks_windows(input_path, output_path)
        else:
            _add_watermarks_basic(input_path, output_path)

        return input_path, True, ""
    except Exception as e:
        return input_path, False, str(e)


def add_watermarks_batch(file_pairs: List[Tuple[str, str]], max_workers: int = None) -> List[Tuple[str, bool, str]]:
    """
    Add watermarks to multiple files in parallel with progress tracking.

    Args:
        file_pairs: List of (input_path, output_path) tuples
        max_workers: Maximum number of parallel workers (default: CPU count)

    Returns:
        List of (input_path, success, error_message) tuples
    """
    if max_workers is None:
        max_workers = min(multiprocessing.cpu_count(), len(file_pairs))

    # Determine method based on availability
    method = "windows" if WINDOWS_AVAILABLE else "basic"

    # Prepare arguments for parallel processing
    args_list = [(input_path, output_path, method) for input_path, output_path in file_pairs]

    results = []

    # Use ProcessPoolExecutor for CPU-bound tasks (Word COM operations)
    with ProcessPoolExecutor(max_workers=max_workers) as executor:
        # Submit all tasks
        future_to_args = {
            executor.submit(process_single_file, args): args
            for args in args_list
        }

        # Process results with progress bar
        with tqdm(total=len(file_pairs), desc="Adding watermarks", unit="file") as pbar:
            for future in as_completed(future_to_args):
                args = future_to_args[future]
                try:
                    result = future.result()
                    results.append(result)

                    input_path, success, error_msg = result
                    if success:
                        pbar.set_postfix_str(f"✅ {os.path.basename(input_path)}")
                    else:
                        pbar.set_postfix_str(f"❌ {os.path.basename(input_path)}: {error_msg}")

                except Exception as exc:
                    input_path = args[0]
                    results.append((input_path, False, str(exc)))
                    pbar.set_postfix_str(f"❌ {os.path.basename(input_path)}: {str(exc)}")

                pbar.update(1)

    return results


def _add_watermarks_windows(input_path: str, output_path: str) -> None:
    """Windows COM-based watermarking (complex, tiled watermarks)."""
    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = win32.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0  # wdAlertsNone

        doc = word.Documents.Open(os.path.abspath(input_path))

        watermark_text = f"{datetime.now().strftime('%Y-%m-%d %H.%M.%S')}_AI Race"

        # Add to each section's primary header so it appears on all pages
        wdHeaderFooterPrimary = getattr(c, "wdHeaderFooterPrimary", 1)
        for section in doc.Sections:
            header = section.Headers(wdHeaderFooterPrimary)
            add_tiled_watermarks_to_header(header, watermark_text)

        # Save output as a new file
        out_abs = os.path.abspath(output_path)
        # Use SaveAs2 for compatibility
        doc.SaveAs2(out_abs)
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


def _add_watermarks_basic(input_path: str, output_path: str) -> None:
    """Basic watermarking for Linux using python-docx (simple text watermark)."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        raise OSError("Basic watermarking requires 'python-docx' library. Install with: pip install python-docx")

    # Load the document
    doc = Document(input_path)

    # Create watermark text
    watermark_text = f"{datetime.now().strftime('%Y-%m-%d %H.%M.%S')}_AI Race"

    # Add watermark to each section
    for section in doc.sections:
        # Get or create header
        header = section.header

        # Create a paragraph for the watermark
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

        # Clear existing content
        for run in paragraph.runs:
            run.clear()

        # Add watermark text
        run = paragraph.add_run(watermark_text)
        run.font.size = Pt(36)
        run.font.color.rgb = RGBColor(200, 200, 200)  # Light gray

        # Center the watermark
        paragraph.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER

        # Make it semi-transparent (if supported)
        try:
            # This might not work in all versions of python-docx
            run.font.color.theme_color = None
            run.font.color.rgb = RGBColor(200, 200, 200)
        except:
            pass

    # Save the document
    doc.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Add complex tiled watermarks to a .docx")
    parser.add_argument(
        "-i",
        "--input",
        default="a.docx",
        help="Path to input .docx (default: a.docx)",
    )
    parser.add_argument(
        "-o",
        "--output",
        default="a_watermarked.docx",
        help="Path to output .docx (default: a_watermarked.docx)",
    )
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    add_watermarks_to_docx(args.input, args.output)


